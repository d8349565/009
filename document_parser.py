"""
文档解析器 - 支持 Markdown/Word/PDF
用于综合报告功能，解析用户提供的初稿文件
"""
import os
import re
import io
import logging
from typing import Dict, Optional


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_file(file_path: str) -> Dict[str, str]:
        """
        解析文档文件，提取标题和内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            {
                'title': '文档标题',
                'content': '完整内容',
                'outline': '提取的大纲',
                'format': 'md' | 'docx' | 'pdf'
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.md':
            return DocumentParser._parse_markdown(file_path)
        elif ext == '.docx':
            return DocumentParser._parse_word(file_path)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    @staticmethod
    def _parse_markdown(file_path: str) -> Dict[str, str]:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题（第一个 # 标题）
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else "未命名文档"
        
        # 提取大纲（所有标题）
        outline_lines = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
        outline = "\n".join([f"{'  ' * (len(level) - 1)}- {text}" for level, text in outline_lines])
        
        return {
            'title': title,
            'content': content,
            'outline': outline,
            'format': 'md'
        }
    
    @staticmethod
    def _parse_word(file_path: str) -> Dict[str, str]:
        """解析Word文档（需要python-docx库）"""
        try:
            import docx
        except ImportError:
            return {
                'title': '需要安装依赖',
                'content': '请运行: pip install python-docx',
                'outline': '',
                'format': 'docx'
            }
        
        doc = docx.Document(file_path)
        
        # 提取标题（第一个段落或第一个标题样式）
        title = "未命名文档"
        for para in doc.paragraphs:
            if para.text.strip():
                title = para.text.strip()
                break
        
        # 提取所有段落内容
        content_lines = []
        outline_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查是否是标题样式
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                content_lines.append(f"{'#' * level} {text}")
                outline_lines.append(f"{'  ' * (level - 1)}- {text}")
            else:
                content_lines.append(text)
        
        content = "\n\n".join(content_lines)
        outline = "\n".join(outline_lines) if outline_lines else "无大纲"
        
        return {
            'title': title,
            'content': content,
            'outline': outline,
            'format': 'docx'
        }
    
    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, str]:
        """解析PDF文档（需要pdfplumber或PyPDF2库）"""
        with open(file_path, 'rb') as f:
            pdf_bytes = f.read()
        return DocumentParser.parse_pdf_from_bytes(pdf_bytes)

    @staticmethod
    def parse_pdf_from_bytes(pdf_bytes: bytes) -> Dict[str, str]:
        """从字节流解析PDF文档内容"""
        logger_names = ["pdfminer", "pdfminer.pdfinterp", "pdfminer.psparser", "pdfminer.converter"]
        previous_levels = {}
        for name in logger_names:
            logger = logging.getLogger(name)
            previous_levels[name] = logger.level
            logger.setLevel(logging.ERROR)
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                # 提取所有文本
                text_pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_pages.append(text)
                
                content = "\n\n".join(text_pages)
                
                # 尝试提取标题（第一页的第一行）
                title = "未命名文档"
                if text_pages:
                    first_lines = text_pages[0].split('\n')
                    for line in first_lines:
                        if line.strip():
                            title = line.strip()
                            break
                
                # PDF大纲提取比较困难，这里简化处理
                outline = "PDF文档大纲提取功能待完善"
                
                return {
                    'title': title,
                    'content': content,
                    'outline': outline,
                    'format': 'pdf'
                }
        except ImportError:
            try:
                import PyPDF2
                
                reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                
                text_pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_pages.append(text)
                
                content = "\n\n".join(text_pages)
                
                title = "未命名文档"
                if text_pages:
                    first_lines = text_pages[0].split('\n')
                    for line in first_lines:
                        if line.strip():
                            title = line.strip()
                            break
                
                outline = "PDF文档大纲提取功能待完善"
                
                return {
                    'title': title,
                    'content': content,
                    'outline': outline,
                    'format': 'pdf'
                }
            except ImportError:
                return {
                    'title': '需要安装依赖',
                    'content': '请运行: pip install pdfplumber 或 pip install PyPDF2',
                    'outline': '',
                    'format': 'pdf'
                }
            except Exception as e:
                return {
                    'title': '解析失败',
                    'content': f'PDF解析出错: {str(e)}',
                    'outline': '',
                    'format': 'pdf'
                }
        except Exception as e:
            return {
                'title': '解析失败',
                'content': f'PDF解析出错: {str(e)}',
                'outline': '',
                'format': 'pdf'
            }
        finally:
            for name, level in previous_levels.items():
                logging.getLogger(name).setLevel(level)


def test_parser():
    """测试文档解析器"""
    print("文档解析器测试")
    print("="*50)
    
    # 测试Markdown
    test_md = """# 测试报告

## 第一章

这是第一章的内容。

### 1.1 小节

这是小节内容。

## 第二章

这是第二章的内容。
"""
    
    # 创建临时测试文件
    test_file = "test_document.md"
    with open(test_file, 'w', encoding='utf-8') as f:
        f.write(test_md)
    
    try:
        result = DocumentParser.parse_file(test_file)
        print(f"标题: {result['title']}")
        print(f"格式: {result['format']}")
        print(f"\n大纲:\n{result['outline']}")
        print(f"\n内容长度: {len(result['content'])} 字符")
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)
    
    print("\n测试完成！")


if __name__ == "__main__":
    test_parser()
